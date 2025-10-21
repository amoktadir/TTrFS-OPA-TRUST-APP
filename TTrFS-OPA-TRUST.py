import streamlit as st
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
from pulp import LpMaximize, LpProblem, LpVariable, lpSum, value
import math

# Set page configuration
st.set_page_config(
    page_title="Integrated MCDM Models",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling
st.markdown("""
<style>
:root {
    --primary: #1f77b4;
    --secondary: #2ca02c;
    --accent: #ff6b6b;
    --background: #f8f9fa;
    --card-bg: #ffffff;
    --text: #262730;
    --border-radius: 12px;
    --shadow: 0 6px 16px rgba(0, 0, 0, 0.08);
}

.main-header {
    font-size: 2.8rem;
    color: var(--primary);
    text-align: center;
    padding: 1.5rem 0;
    font-weight: 700;
    margin-bottom: 0.5rem;
    background: linear-gradient(135deg, var(--primary), var(--secondary));
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.section-header {
    font-size: 1.6rem;
    color: var(--primary);
    border-left: 5px solid var(--secondary);
    padding-left: 1rem;
    margin: 2rem 0 1.5rem 0;
    font-weight: 600;
}

.panel {
    background-color: var(--card-bg);
    border-radius: var(--border-radius);
    padding: 1.8rem;
    box-shadow: var(--shadow);
    margin-bottom: 1.5rem;
    border: 1px solid #e0e0e0;
    transition: transform 0.2s ease, box-shadow 0.2s ease;
}

.panel:hover {
    transform: translateY(-3px);
    box-shadow: 0 8px 20px rgba(0, 0, 0, 0.12);
}

.metric-card {
    background: linear-gradient(135deg, var(--card-bg), #f7f9fc);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--shadow);
    text-align: center;
    border: 1px solid #e8e8e8;
    height: 100%;
}

.metric-value {
    font-size: 2rem;
    font-weight: 700;
    color: var(--primary);
    margin: 0.5rem 0;
}

.metric-label {
    font-size: 0.9rem;
    color: #666;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

.result-table {
    background-color: var(--card-bg);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--shadow);
    margin: 1.5rem 0;
}

.optimization-formulation {
    background-color: #f8f9fa;
    border-left: 4px solid var(--primary);
    padding: 1.5rem;
    border-radius: 8px;
    margin: 1.5rem 0;
    font-family: 'Courier New', monospace;
    font-size: 0.9rem;
    line-height: 1.6;
}

.optimization-section {
    background-color: #fff;
    border: 1px solid #e0e0e0;
    border-radius: 8px;
    padding: 1.5rem;
    margin: 1.5rem 0;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}

.stButton>button {
    background: linear-gradient(135deg, var(--primary), var(--secondary));
    color: white;
    border: none;
    padding: 0.8rem 2rem;
    border-radius: 50px;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: 0 4px 6px rgba(50, 50, 93, 0.11), 0 1px 3px rgba(0, 0, 0, 0.08);
    width: 100%;
    margin-top: 1rem;
}

.stButton>button:hover {
    transform: translateY(-2px);
    box-shadow: 0 7px 14px rgba(50, 50, 93, 0.1), 0 3px 6px rgba(0, 0, 0, 0.08);
    background: linear-gradient(135deg, var(--secondary), var(--primary));
}

.criteria-input {
    padding: 0.8rem;
    border-radius: 8px;
    border: 1px solid #e0e0e0;
    margin-bottom: 0.8rem;
    transition: border 0.3s ease;
    width: 100%;
}

.criteria-input:focus {
    border-color: var(--primary);
    outline: none;
    box-shadow: 0 0 0 2px rgba(31, 119, 180, 0.2);
}

.instruction-box {
    background-color: #f0f7ff;
    border-left: 4px solid var(--primary);
    padding: 1rem 1.5rem;
    border-radius: 4px;
    margin: 1rem 0;
    font-size: 0.95rem;
}

.success-box {
    background-color: #f0fff4;
    border-left: 4px solid var(--secondary);
    padding: 1rem 1.5rem;
    border-radius: 4px;
    margin: 1rem 0;
}

.warning-box {
    background-color: #fffaf0;
    border-left: 4px solid #ffb347;
    padding: 1rem 1.5rem;
    border-radius: 4px;
    margin: 1rem 0;
}

.error-box {
    background-color: #fff5f5;
    border-left: 4px solid #ff6b6b;
    padding: 1rem 1.5rem;
    border-radius: 4px;
    margin: 1rem 0;
}

.comparison-scale {
    background-color: #f8f9fa;
    padding: 1rem;
    border-radius: 8px;
    margin: 1rem 0;
    font-size: 0.9rem;
}

.scale-item {
    display: flex;
    justify-content: space-between;
    padding: 0.3rem 0;
    border-bottom: 1px dashed #e0e0e0;
}

.scale-item:last-child {
    border-bottom: none;
}

.footer {
    text-align: center;
    margin-top: 3rem;
    padding: 1.5rem;
    color: #666;
    font-size: 0.9rem;
    border-top: 1px solid #eaeaea;
}

.logo-container {
    text-align: center;
    margin-bottom: 1.5rem;
}

.logo {
    font-size: 3rem;
    margin-bottom: 0.5rem;
}

.stSlider .st-emotion-cache-1adqxcy {
    background: linear-gradient(to right, var(--primary), var(--secondary));
}

.stSelectbox div div div {
    color: var(--text);
}

input[type="number"] {
    border: 1px solid #e0e0e0;
    border-radius: 8px;
    padding: 0.5rem;
}

</style>
""", unsafe_allow_html=True)

# ==================== OPA MODEL FUNCTIONS ====================

# Linguistic terms to TrFN for OPA
ling_to_trfn_opa = {
    'ELI': (1.0, 1.5, 2.5, 3.0),
    'VLI': (2.0, 2.5, 3.5, 4.0),
    'LI': (3.0, 3.5, 4.5, 5.0),
    'MI': (4.0, 4.5, 5.5, 6.0),
    'HI': (5.0, 5.5, 6.5, 7.0),
    'VHI': (6.0, 6.5, 7.5, 8.0),
    'EHI': (7.0, 7.5, 8.5, 9.0),
}

linguistic_options_opa = list(ling_to_trfn_opa.keys())

def trig_geom_component(values, weights):
    n = len(values)
    s = sum(values)
    if s == 0:
        return 0
    prod = 1
    for v, w in zip(values, weights):
        f = v / s
        prod *= np.sin(np.pi * f / 2) ** w
    arc = np.arcsin(prod)
    return s * (2 / np.pi) * arc

def aggregate_ftwg(trfn_list, weights):
    ls = [tf[0] for tf in trfn_list]
    ms = [tf[1] for tf in trfn_list]
    us = [tf[2] for tf in trfn_list]
    ws = [tf[3] for tf in trfn_list]
    
    l = trig_geom_component(ls, weights)
    m = trig_geom_component(ms, weights)
    u = trig_geom_component(us, weights)
    ww = trig_geom_component(ws, weights)
    
    return (l, m, u, ww)

def defuzz(trf):
    return (2 * trf[0] + 7 * trf[1] + 7 * trf[2] + 2 * trf[3]) / 18

def solve_fuzzy_opa(coeff_list, n):
    prob = LpProblem("Trapezoidal_Fuzzy_Trig_OPA", LpMaximize)
    
    # Define decision variables for the weights (w)
    w_l = [LpVariable(f"w_l_{i}", lowBound=0) for i in range(n)]
    w_m = [LpVariable(f"w_m_{i}", lowBound=0) for i in range(n)]
    w_u = [LpVariable(f"w_u_{i}", lowBound=0) for i in range(n)]
    w_w = [LpVariable(f"w_w_{i}", lowBound=0) for i in range(n)]
    
    # Define auxiliary variables for the objective function (Psi)
    Psi_l = LpVariable("Psi_l", lowBound=0)
    Psi_m = LpVariable("Psi_m", lowBound=0)
    Psi_u = LpVariable("Psi_u", lowBound=0)
    Psi_w = LpVariable("Psi_w", lowBound=0)
    
    # Objective Function: Maximize the defuzzified value of Psi
    prob += (2 * Psi_l + 7 * Psi_m + 7 * Psi_u + 2 * Psi_w) / 18
    
    # Constraints
    for i in range(n):
        prob += w_l[i] <= w_m[i]
        prob += w_m[i] <= w_u[i]
        prob += w_u[i] <= w_w[i]
    
    # Normalization constraints
    prob += lpSum(w_l) == 0.8
    prob += lpSum(w_m) == 0.9
    prob += lpSum(w_u) == 1.1
    prob += lpSum(w_w) == 1.2
    
    # Core OPA constraints comparing adjacent ranked criteria
    for a in range(n - 1):
        prob += coeff_list[a][0] * (w_l[a] - w_w[a + 1]) >= Psi_l
        prob += coeff_list[a][1] * (w_m[a] - w_u[a + 1]) >= Psi_m
        prob += coeff_list[a][2] * (w_u[a] - w_m[a + 1]) >= Psi_u
        prob += coeff_list[a][3] * (w_w[a] - w_l[a + 1]) >= Psi_w
    
    # OPA constraints for the last-ranked criterion
    prob += coeff_list[n - 1][0] * w_l[n - 1] >= Psi_l
    prob += coeff_list[n - 1][1] * w_m[n - 1] >= Psi_m
    prob += coeff_list[n - 1][2] * w_u[n - 1] >= Psi_u
    prob += coeff_list[n - 1][3] * w_w[n - 1] >= Psi_w
    
    # Solve the problem
    status = prob.solve()
    
    if status != 1:
        st.error("Optimization failed. The problem may be infeasible. Please check your inputs.")
        return None, None
    
    # Extract and clean results
    weights = []
    for i in range(n):
        wl = max(0, value(w_l[i]))
        wm = max(0, value(w_m[i]))
        wu = max(0, value(w_u[i]))
        ww = max(0, value(w_w[i]))
        weights.append((wl, wm, wu, ww))
    
    psi = (
        max(0, value(Psi_l)),
        max(0, value(Psi_m)),
        max(0, value(Psi_u)),
        max(0, value(Psi_w))
    )
    
    return weights, psi

def create_opa_word_document(criteria, theta, defuzz_values, coeff, ranked_criteria, weights, psi, num_experts, expert_weights):
    doc = Document()
    title = doc.add_heading('Trigonometric Trapezoidal Fuzzy OPA Results - Multiple Experts', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'This document contains the results of the Fuzzy Trigonometric OPA analysis for {num_experts} experts.')
    doc.add_paragraph('')
    
    # Expert Weights
    doc.add_heading('Expert Weights', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Expert'
    hdr_cells[1].text = 'Weight'
    
    for e in range(num_experts):
        row_cells = table.add_row().cells
        row_cells[0].text = f"Expert {e+1}"
        row_cells[1].text = f"{expert_weights[e]:.4f}"
    
    doc.add_paragraph('')
    
    # Aggregated Theta
    doc.add_heading('Aggregated Fuzzy Importance (Theta)', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Criterion'
    hdr_cells[1].text = 'l'
    hdr_cells[2].text = 'm'
    hdr_cells[3].text = 'u'
    hdr_cells[4].text = 'w'
    hdr_cells[5].text = 'Defuzzified'
    
    for i, crit in enumerate(criteria):
        row_cells = table.add_row().cells
        row_cells[0].text = crit
        row_cells[1].text = f"{theta[i][0]:.4f}"
        row_cells[2].text = f"{theta[i][1]:.4f}"
        row_cells[3].text = f"{theta[i][2]:.4f}"
        row_cells[4].text = f"{theta[i][3]:.4f}"
        row_cells[5].text = f"{defuzz_values[i]:.4f}"
    
    doc.add_paragraph('')
    
    # Coefficients
    doc.add_heading('Coefficients for Fuzzy OPA', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Criterion'
    hdr_cells[1].text = 'Coeff l'
    hdr_cells[2].text = 'Coeff m'
    hdr_cells[3].text = 'Coeff u'
    hdr_cells[4].text = 'Coeff w'
    
    for i, crit in enumerate(criteria):
        row_cells = table.add_row().cells
        row_cells[0].text = crit
        row_cells[1].text = f"{coeff[i][0]:.4f}"
        row_cells[2].text = f"{coeff[i][1]:.4f}"
        row_cells[3].text = f"{coeff[i][2]:.4f}"
        row_cells[4].text = f"{coeff[i][3]:.4f}"
    
    doc.add_paragraph('')
    
    # Ranked Criteria
    doc.add_heading('Ranked Criteria', level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rank'
    hdr_cells[1].text = 'Criterion'
    hdr_cells[2].text = 'Weight l'
    hdr_cells[3].text = 'Weight m'
    hdr_cells[4].text = 'Weight u'
    hdr_cells[5].text = 'Weight w'
    hdr_cells[6].text = 'Defuzzified'
    
    for rank, (crit_idx, w) in enumerate(ranked_criteria):
        row_cells = table.add_row().cells
        row_cells[0].text = str(rank + 1)
        row_cells[1].text = criteria[crit_idx]
        row_cells[2].text = f"{w[0]:.4f}"
        row_cells[3].text = f"{w[1]:.4f}"
        row_cells[4].text = f"{w[2]:.4f}"
        row_cells[5].text = f"{w[3]:.4f}"
        row_cells[6].text = f"{defuzz(w):.4f}"
    
    doc.add_paragraph('')
    
    # Psi
    doc.add_heading('Psi Value', level=1)
    doc.add_paragraph(f'Psi: ({psi[0]:.4f}, {psi[1]:.4f}, {psi[2]:.4f}, {psi[3]:.4f})')
    doc.add_paragraph(f'Defuzzified Psi: {defuzz(psi):.4f}')
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def display_optimization_formulation(coeff_list, n, criteria_names):
    """Display the optimization problem formulation"""
    st.markdown('<div class="optimization-section">', unsafe_allow_html=True)
    st.markdown('<h3 class="section-header">TTrFS-OPA Optimization Problem Formulation</h3>', unsafe_allow_html=True)
    
    # Objective Function
    st.markdown("""
    <div class="optimization-formulation">
    <strong>Objective Function:</strong><br>
    Maximize: (2·Ψ_l + 7·Ψ_m + 7·Ψ_u + 2·Ψ_w) / 18
    </div>
    """, unsafe_allow_html=True)
    
    # Decision Variables
    st.markdown("""
    <div class="optimization-formulation">
    <strong>Decision Variables:</strong><br>
    w_lᵢ, w_mᵢ, w_uᵢ, w_wᵢ ≥ 0 for i = 1,...,n<br>
    Ψ_l, Ψ_m, Ψ_u, Ψ_w ≥ 0
    </div>
    """, unsafe_allow_html=True)
    
    # Constraints
    st.markdown("""
    <div class="optimization-formulation">
    <strong>Constraints:</strong><br>
    1. Fuzzy weight ordering:<br>
    &nbsp;&nbsp;w_lᵢ ≤ w_mᵢ ≤ w_uᵢ ≤ w_wᵢ for all i<br><br>
    
    2. Normalization constraints:<br>
    &nbsp;&nbsp;∑w_lᵢ = 0.8<br>
    &nbsp;&nbsp;∑w_mᵢ = 0.9<br>
    &nbsp;&nbsp;∑w_uᵢ = 1.1<br>
    &nbsp;&nbsp;∑w_wᵢ = 1.2<br><br>
    
    3. Adjacent criteria constraints (for a = 1 to n-1):<br>
    &nbsp;&nbsp;θₗᵃ·(w_lᵃ - w_wᵃ⁺¹) ≥ Ψ_l<br>
    &nbsp;&nbsp;θₘᵃ·(w_mᵃ - w_uᵃ⁺¹) ≥ Ψ_m<br>
    &nbsp;&nbsp;θᵤᵃ·(w_uᵃ - w_mᵃ⁺¹) ≥ Ψ_u<br>
    &nbsp;&nbsp;θ_wᵃ·(w_wᵃ - w_lᵃ⁺¹) ≥ Ψ_w<br><br>
    
    4. Last criterion constraint:<br>
    &nbsp;&nbsp;θₗⁿ·w_lⁿ ≥ Ψ_l<br>
    &nbsp;&nbsp;θₘⁿ·w_mⁿ ≥ Ψ_m<br>
    &nbsp;&nbsp;θᵤⁿ·w_uⁿ ≥ Ψ_u<br>
    &nbsp;&nbsp;θ_wⁿ·w_wⁿ ≥ Ψ_w
    </div>
    """, unsafe_allow_html=True)
    
    # Display specific coefficients for this problem
    st.markdown("""
    <div class="optimization-formulation">
    <strong>Problem-specific Coefficients (θ):</strong>
    </div>
    """, unsafe_allow_html=True)
    
    coeff_df = pd.DataFrame({
        'Criterion': criteria_names,
        'θ_l': [f"{coeff[0]:.4f}" for coeff in coeff_list],
        'θ_m': [f"{coeff[1]:.4f}" for coeff in coeff_list],
        'θ_u': [f"{coeff[2]:.4f}" for coeff in coeff_list],
        'θ_w': [f"{coeff[3]:.4f}" for coeff in coeff_list]
    })
    st.dataframe(coeff_df, use_container_width=True, hide_index=True)
    
    # Mathematical notation explanation
    with st.expander("Mathematical Notation Explanation"):
        st.markdown("""
        **Notation:**
        - **w_lᵢ, w_mᵢ, w_uᵢ, w_wᵢ**: Trapezoidal fuzzy weight components for criterion i
        - **Ψ_l, Ψ_m, Ψ_u, Ψ_w**: Auxiliary variables representing satisfaction degrees
        - **θₗᵢ, θₘᵢ, θᵤᵢ, θ_wᵢ**: Coefficient components for criterion i
        - **n**: Number of criteria
        
        **Coefficient Calculation:**
        θₗᵢ = min_l / wᵢ, θₘᵢ = min_l / uᵢ, θᵤᵢ = min_l / mᵢ, θ_wᵢ = min_l / lᵢ
        where min_l is the minimum l-component across all fuzzy importance values
        """)
    
    st.markdown('</div>', unsafe_allow_html=True)

def opa_model():
    st.markdown('<div class="logo-container">', unsafe_allow_html=True)
    st.markdown('<div class="logo">⚖️</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="main-header">Trigonometric Trapezoidal Fuzzy OPA Analysis with Multiple Experts</h1>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("""
    <div style="text-align: center; margin-bottom: 2.5rem; color: #555;">
    <p style="font-size: 1.1rem;">This application implements the Trigonometric Trapezoidal Fuzzy Ordinal Priority Approach (OPA) for multi-criteria decision-making with multiple experts.
    </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize session state for OPA
    if 'opa_criteria' not in st.session_state:
        st.session_state.opa_criteria = [f"Criterion {i+1}" for i in range(5)]
    if 'opa_num_criteria' not in st.session_state:
        st.session_state.opa_num_criteria = 5
    if 'opa_num_experts' not in st.session_state:
        st.session_state.opa_num_experts = 2
    if 'opa_results_calculated' not in st.session_state:
        st.session_state.opa_results_calculated = False
    if 'opa_expert_data' not in st.session_state:
        st.session_state.opa_expert_data = {}
    if 'opa_expert_weights' not in st.session_state:
        st.session_state.opa_expert_weights = [1.0 / st.session_state.opa_num_experts] * st.session_state.opa_num_experts
    
    # Create two columns for layout
    left_col, right_col = st.columns([1, 1])
    
    with left_col:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.markdown('<h2 class="section-header">Step 1: Define Decision Criteria and Experts</h2>', unsafe_allow_html=True)
        
        num_experts = st.number_input(
            "Number of experts",
            min_value=1,
            max_value=15,
            value=st.session_state.opa_num_experts,
            step=1,
            key="opa_num_experts_input",
            help="Select between 1 to 15 experts"
        )
        
        num_criteria = st.number_input(
            "Number of criteria",
            min_value=3,
            max_value=25,
            value=st.session_state.opa_num_criteria,
            step=1,
            key="opa_num_criteria_input",
            help="Select between 3 to 25 criteria"
        )
        
        if num_criteria != st.session_state.opa_num_criteria or num_experts != st.session_state.opa_num_experts:
            st.session_state.opa_num_criteria = num_criteria
            st.session_state.opa_num_experts = num_experts
            if len(st.session_state.opa_criteria) < num_criteria:
                for i in range(len(st.session_state.opa_criteria), num_criteria):
                    st.session_state.opa_criteria.append(f"Criterion {i+1}")
            else:
                st.session_state.opa_criteria = st.session_state.opa_criteria[:num_criteria]
            st.session_state.opa_expert_data = {}
            st.session_state.opa_expert_weights = [1.0 / num_experts] * num_experts
            st.session_state.opa_results_calculated = False
        
        criteria = []
        for i in range(num_criteria):
            criterion = st.text_input(
                f"Criterion {i+1}",
                value=st.session_state.opa_criteria[i],
                key=f"opa_criterion_{i}",
                placeholder=f"Enter name for criterion {i+1}"
            )
            criteria.append(criterion)
        
        st.session_state.opa_criteria = criteria
        
        st.markdown('<h3>Expert Weights</h3>', unsafe_allow_html=True)
        st.markdown("""
        <div class="instruction-box">
        <strong>Important:</strong> The sum of all expert weights must equal exactly 1.00
        </div>
        """, unsafe_allow_html=True)
        
        expert_weights = []
        for e in range(num_experts):
            w = st.number_input(
                f"Weight for Expert {e+1}",
                min_value=0.0,
                max_value=1.0,
                value=st.session_state.opa_expert_weights[e],
                step=0.01,
                format="%.2f",
                key=f"opa_expert_w_{e}"
            )
            expert_weights.append(w)
        
        # Calculate and display sum of weights
        sum_w = sum(expert_weights)
        st.markdown(f"**Sum of expert weights: {sum_w:.2f}**")
        
        # Validate weights sum
        if abs(sum_w - 1.0) > 0.01:
            st.error(f"❌ Sum of expert weights must equal 1.00. Current sum: {sum_w:.2f}")
            st.session_state.opa_weights_valid = False
        else:
            # Normalize to ensure exact sum of 1.0
            if sum_w > 0:
                normalized_weights = [w / sum_w for w in expert_weights]
                st.success(f"✅ Weights are valid! Sum: {sum(normalized_weights):.2f}")
            else:
                normalized_weights = [1.0 / num_experts] * num_experts
                st.error("❌ Weights cannot be all zero!")
            st.session_state.opa_expert_weights = normalized_weights
            st.session_state.opa_weights_valid = True
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with right_col:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.markdown('<h2 class="section-header">Expert Assessments</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="instruction-box">
        <strong>Linguistic Terms:</strong><br>
        ELI (Extremely Low), VLI (Very Low), LI (Low), MI (Medium), HI (High), VHI (Very High), EHI (Extremely High)
        </div>
        """, unsafe_allow_html=True)
        
        # Individual expert tabs with data editor only
        tabs = st.tabs([f"Expert {e+1}" for e in range(num_experts)])
        
        for e, tab in enumerate(tabs):
            with tab:
                if f'expert_{e}' not in st.session_state.opa_expert_data:
                    st.session_state.opa_expert_data[f'expert_{e}'] = ['MI'] * num_criteria
                
                st.markdown(f'<h3>Expert {e+1} Input</h3>', unsafe_allow_html=True)
                
                # Data Editor for quick input
                st.markdown("**Edit ratings directly in the table below:**")
                
                # Create DataFrame for data editor
                df_data = {
                    'Criterion': criteria,
                    'Rating': st.session_state.opa_expert_data[f'expert_{e}']
                }
                df = pd.DataFrame(df_data)
                
                # Configure the data editor
                edited_df = st.data_editor(
                    df,
                    column_config={
                        "Rating": st.column_config.SelectboxColumn(
                            "Rating",
                            options=linguistic_options_opa,
                            required=True,
                            width="medium"
                        )
                    },
                    use_container_width=True,
                    key=f"opa_data_editor_expert_{e}"
                )
                
                # Update session state with edited data
                if not edited_df.equals(df):
                    st.session_state.opa_expert_data[f'expert_{e}'] = edited_df['Rating'].tolist()
                    st.success(f"Expert {e+1} ratings updated!")
                    st.rerun()
        
        # Calculate button with validation
        st.markdown('<div style="text-align: center; margin: 2.5rem 0;">', unsafe_allow_html=True)
        
        # Check if all conditions are met for calculation
        all_ratings_available = all(f'expert_{e}' in st.session_state.opa_expert_data for e in range(num_experts))
        weights_valid = getattr(st.session_state, 'opa_weights_valid', False)
        
        if not all_ratings_available:
            st.error("❌ Please provide ratings for all experts before calculating weights.")
            calculate_disabled = True
        elif not weights_valid:
            st.error("❌ Please fix expert weights (sum must equal 1.00) before calculating.")
            calculate_disabled = True
        else:
            calculate_disabled = False
        
        if st.button("Calculate Weights", key="opa_calculate_button", use_container_width=True, disabled=calculate_disabled):
            with st.spinner("Calculating weights..."):
                # Final validation before calculation
                if not all_ratings_available:
                    st.error("Please provide ratings for all experts before calculating weights.")
                elif not weights_valid:
                    st.error("Expert weights sum must equal 1.00. Please adjust the weights.")
                else:
                    # Aggregate theta using FTWG
                    theta = []
                    defuzz_values = []
                    for j in range(num_criteria):
                        trfn_list = [ling_to_trfn_opa[st.session_state.opa_expert_data[f'expert_{e}'][j]] for e in range(num_experts)]
                        aggregated_trf = aggregate_ftwg(trfn_list, st.session_state.opa_expert_weights)
                        theta.append(aggregated_trf)
                        defuzz_values.append(defuzz(aggregated_trf))
                    
                    # Compute min_l
                    min_l = min(t[0] for t in theta if t[0] > 0) if any(t[0] > 0 for t in theta) else 1.0
                    
                    # Compute coefficients
                    coeff = []
                    for t in theta:
                        if t[0] == 0 or t[1] == 0 or t[2] == 0 or t[3] == 0:
                            c = (0, 0, 0, 0)
                        else:
                            c = (min_l / t[3], min_l / t[2], min_l / t[1], min_l / t[0])
                        coeff.append(c)
                    
                    # Rank criteria based on defuzz(theta)
                    sorted_indices = np.argsort(defuzz_values)[::-1]
                    
                    # Sort coeff for OPA
                    coeff_sorted = [coeff[idx] for idx in sorted_indices]
                    
                    # Solve LP
                    weights_sorted, psi = solve_fuzzy_opa(coeff_sorted, num_criteria)
                    
                    if weights_sorted is not None:
                        # Map weights back to original order
                        weights = [None] * num_criteria
                        for rank, idx in enumerate(sorted_indices):
                            weights[idx] = weights_sorted[rank]
                        
                        # Ranked criteria with weights
                        ranked_criteria = [(sorted_indices[k], weights_sorted[k]) for k in range(num_criteria)]
                        
                        st.session_state.opa_theta = theta
                        st.session_state.opa_defuzz_values = defuzz_values
                        st.session_state.opa_coeff = coeff
                        st.session_state.opa_ranked_criteria = ranked_criteria
                        st.session_state.opa_weights = weights
                        st.session_state.opa_psi = psi
                        st.session_state.opa_results_calculated = True
                        st.success("Calculation completed successfully!")
        
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Results section
    if st.session_state.opa_results_calculated:
        st.markdown('<h2 class="section-header">Results</h2>', unsafe_allow_html=True)
        
        # Expert Weights
        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader("Normalized Expert Weights")
        df_expert_w = pd.DataFrame({
            'Expert': [f"Expert {e+1}" for e in range(st.session_state.opa_num_experts)],
            'Weight': st.session_state.opa_expert_weights
        })
        st.dataframe(df_expert_w, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Aggregated Theta
        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader("Aggregated Fuzzy Importance (Theta)")
        df_theta = pd.DataFrame({
            'Criterion': st.session_state.opa_criteria,
            'l': [t[0] for t in st.session_state.opa_theta],
            'm': [t[1] for t in st.session_state.opa_theta],
            'u': [t[2] for t in st.session_state.opa_theta],
            'w': [t[3] for t in st.session_state.opa_theta],
            'Defuzzified': st.session_state.opa_defuzz_values
        })
        st.dataframe(df_theta, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Coefficients
        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader("Coefficients for Fuzzy OPA")
        df_coeff = pd.DataFrame({
            'Criterion': st.session_state.opa_criteria,
            'Coeff l': [c[0] for c in st.session_state.opa_coeff],
            'Coeff m': [c[1] for c in st.session_state.opa_coeff],
            'Coeff u': [c[2] for c in st.session_state.opa_coeff],
            'Coeff w': [c[3] for c in st.session_state.opa_coeff]
        })
        st.dataframe(df_coeff, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # NEW: Optimization Problem Formulation
        if hasattr(st.session_state, 'opa_coeff') and hasattr(st.session_state, 'opa_criteria'):
            # Get the sorted coefficients for the optimization problem
            sorted_indices = np.argsort(st.session_state.opa_defuzz_values)[::-1]
            coeff_sorted = [st.session_state.opa_coeff[idx] for idx in sorted_indices]
            criteria_sorted = [st.session_state.opa_criteria[idx] for idx in sorted_indices]
            
            display_optimization_formulation(
                coeff_sorted, 
                st.session_state.opa_num_criteria, 
                criteria_sorted
            )
        
        # Weights
        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader("Fuzzy Weights")
        df_weights = pd.DataFrame({
            'Criterion': st.session_state.opa_criteria,
            'l': [w[0] for w in st.session_state.opa_weights],
            'm': [w[1] for w in st.session_state.opa_weights],
            'u': [w[2] for w in st.session_state.opa_weights],
            'w': [w[3] for w in st.session_state.opa_weights],
            'Defuzzified': [defuzz(w) for w in st.session_state.opa_weights]
        })
        st.dataframe(df_weights, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Ranked
        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader("Ranked Criteria and Weights")
        df_ranked = pd.DataFrame({
            'Rank': list(range(1, st.session_state.opa_num_criteria + 1)),
            'Criterion': [st.session_state.opa_criteria[idx] for idx, w in st.session_state.opa_ranked_criteria],
            'Weight l': [w[0] for idx, w in st.session_state.opa_ranked_criteria],
            'Weight m': [w[1] for idx, w in st.session_state.opa_ranked_criteria],
            'Weight u': [w[2] for idx, w in st.session_state.opa_ranked_criteria],
            'Weight w': [w[3] for idx, w in st.session_state.opa_ranked_criteria],
            'Defuzzified': [defuzz(w) for idx, w in st.session_state.opa_ranked_criteria]
        })
        st.dataframe(df_ranked, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Psi
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.markdown('<div class="metric-label">Psi (l, m, u, w)</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="metric-value">{st.session_state.opa_psi[0]:.4f}, {st.session_state.opa_psi[1]:.4f}, {st.session_state.opa_psi[2]:.4f}, {st.session_state.opa_psi[3]:.4f}</div>', unsafe_allow_html=True)
        st.markdown('<div class="metric-label">Defuzzified Psi</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="metric-value">{defuzz(st.session_state.opa_psi):.4f}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Export button
        st.markdown('<div style="text-align: center; margin: 2.5rem 0;">', unsafe_allow_html=True)
        doc_bytes = create_opa_word_document(
            st.session_state.opa_criteria,
            st.session_state.opa_theta,
            st.session_state.opa_defuzz_values,
            st.session_state.opa_coeff,
            st.session_state.opa_ranked_criteria,
            st.session_state.opa_weights,
            st.session_state.opa_psi,
            st.session_state.opa_num_experts,
            st.session_state.opa_expert_weights
        )
        
        st.download_button(
            label="Export Results to Word",
            data=doc_bytes,
            file_name="Trigonometric_Fuzzy_OPA_Results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    with st.expander("Learn more about the Trigonometric Trapezoidal Fuzzy OPA (TTrF-OPA) Method"):
        st.markdown("""
        **Steps of Trigonometric Trapezoidal Fuzzy OPA (TTrF-OPA):**
        
        1. Experts provide linguistic assessments for the criteria importance.
        2. Convert to trapezoidal fuzzy numbers and aggregate using trigonometric trapezoidal fuzzy weighted geometric (TTrFWG) operator with expert weights.
        3. Compute weight-coefficients using $min_l / w$, $min_l / u$, $min_l / m$, $min_l / l$.
        4. Defuzzify to rank criteria to form contraints.
        5. Formulate and solve a fuzzy linear programming model to find the final weights and defuzzify and rank.
        
        **Linguistic Terms and TrFN:**
        ELI (Extremely Low Importance), VLI (Very Low Importance), LI (Low Importance), MI (Medium Importance), HI (High Importance), VHI (Very High Importance), EHI (Extremely High Importance)
        - ELI: (1.0, 1.5, 2.5, 3.0)
        - VLI: (2.0, 2.5, 3.5, 4.0)
        - LI: (3.0, 3.5, 4.5, 5.0)
        - MI: (4.0, 4.5, 5.5, 6.0)
        - HI: (5.0, 5.5, 6.5, 7.0)
        - VHI: (6.0, 6.5, 7.5, 8.0)
        - EHI: (7.0, 7.5, 8.5, 9.0)
        """)

# ==================== TRUST MODEL FUNCTIONS ====================

# Linguistic terms to TrFS mapping for TTrFS-TRUST
ling_to_trfn_trust = {
    'ELI': (0.00, 0.10, 0.20, 0.30),
    'VLI': (0.10, 0.20, 0.30, 0.40),
    'LI': (0.20, 0.30, 0.40, 0.50),
    'MLI': (0.30, 0.40, 0.50, 0.60),
    'MI': (0.40, 0.50, 0.60, 0.70),
    'MHI': (0.50, 0.60, 0.70, 0.80),
    'HI': (0.60, 0.70, 0.80, 0.90),
    'VHI': (0.70, 0.80, 0.90, 1.00),
    'EHI': (0.80, 0.90, 1.00, 1.00),
}

linguistic_options_trust = list(ling_to_trfn_trust.keys())

def defuzzify_trfn(trf):
    return (2 * trf[0] + 7 * trf[1] + 7 * trf[2] + 2 * trf[3]) / 18

def crisp_to_trfs(x, alpha=0.05):
    return (x, x + alpha, x + 2*alpha, x + 3*alpha)

def create_trust_word_document(all_data):
    doc = Document()
    title = doc.add_heading('TRUST Method Results - Multi-normalization Multi-distance Assessment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Problem Setup
    doc.add_heading('Problem Setup', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    setup_data = [
        ('Number of Alternatives', all_data['n_alternatives']),
        ('Number of Criteria', all_data['n_criteria']),
        ('Number of Experts', all_data['n_experts']),
        ('Alpha Parameters', f"∂₁: {all_data['alpha'][0]}, ∂₂: {all_data['alpha'][1]}, ∂₃: {all_data['alpha'][2]}, ∂₄: {all_data['alpha'][3]}"),
        ('Beta Parameter', all_data['beta'])
    ]
    
    for i, (label, value) in enumerate(setup_data):
        row_cells = table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
    
    doc.add_paragraph('')
    
    # Decision Matrix
    doc.add_heading('Decision Matrix', level=1)
    decision_matrix = all_data['decision_matrix']
    alternatives = all_data['alternatives']
    criteria = all_data['criteria']
    
    table = doc.add_table(rows=len(alternatives)+1, cols=len(criteria)+1)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alternative'
    for j, criterion in enumerate(criteria):
        hdr_cells[j+1].text = criterion
    
    # Data
    for i, alternative in enumerate(alternatives):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = alternative
        for j, criterion in enumerate(criteria):
            row_cells[j+1].text = f"{decision_matrix[i, j]:.4f}"
    
    doc.add_paragraph('')
    
    # Normalization Matrices
    doc.add_heading('Normalization Matrices', level=1)
    
    # Linear Ratio-based
    doc.add_heading('Linear Ratio-based Normalization (r_ij)', level=2)
    table = doc.add_table(rows=len(alternatives)+1, cols=len(criteria)+1)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alternative'
    for j, criterion in enumerate(criteria):
        hdr_cells[j+1].text = criterion
    
    for i, alternative in enumerate(alternatives):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = alternative
        for j, criterion in enumerate(criteria):
            row_cells[j+1].text = f"{all_data['r_matrix'][i, j]:.4f}"
    
    doc.add_paragraph('')
    
    # Linear Sum-based
    doc.add_heading('Linear Sum-based Normalization (s_ij)', level=2)
    table = doc.add_table(rows=len(alternatives)+1, cols=len(criteria)+1)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alternative'
    for j, criterion in enumerate(criteria):
        hdr_cells[j+1].text = criterion
    
    for i, alternative in enumerate(alternatives):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = alternative
        for j, criterion in enumerate(criteria):
            row_cells[j+1].text = f"{all_data['s_matrix'][i, j]:.4f}"
    
    doc.add_paragraph('')
    
    # Max-Min Normalization
    doc.add_heading('Max-Min Normalization (m_ij)', level=2)
    table = doc.add_table(rows=len(alternatives)+1, cols=len(criteria)+1)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alternative'
    for j, criterion in enumerate(criteria):
        hdr_cells[j+1].text = criterion
    
    for i, alternative in enumerate(alternatives):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = alternative
        for j, criterion in enumerate(criteria):
            row_cells[j+1].text = f"{all_data['m_matrix'][i, j]:.4f}"
    
    doc.add_paragraph('')
    
    # Aggregated Normalized Matrix
    doc.add_heading('Aggregated Normalized Matrix (h_ij)', level=1)
    table = doc.add_table(rows=len(alternatives)+1, cols=len(criteria)+1)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alternative'
    for j, criterion in enumerate(criteria):
        hdr_cells[j+1].text = criterion
    
    for i, alternative in enumerate(alternatives):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = alternative
        for j, criterion in enumerate(criteria):
            row_cells[j+1].text = f"{all_data['h_matrix'][i, j]:.4f}"
    
    doc.add_paragraph('')
    
    # Final Results
    doc.add_heading('Final Ranking Results', level=1)
    table = doc.add_table(rows=len(all_data['final_results'])+1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rank'
    hdr_cells[1].text = 'Alternative'
    hdr_cells[2].text = '∑℘_ik'
    hdr_cells[3].text = '∑ℌ_ik'
    hdr_cells[4].text = 'ℒ Score'
    
    for i, row in all_data['final_results'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Rank'])
        row_cells[1].text = row['Alternative']
        row_cells[2].text = f"{row['∑℘_ik']:.4f}"
        row_cells[3].text = f"{row['∑ℌ_ik']:.4f}"
        row_cells[4].text = f"{row['ℒ Score']:.4f}"
    
    doc.add_paragraph('')
    doc.add_paragraph(f"**Best Alternative: {all_data['best_alternative']}** with score: {all_data['best_score']:.4f}")
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def trust_model():
    st.markdown('<div class="logo-container">', unsafe_allow_html=True)
    st.markdown('<div class="logo">📊</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="main-header">TTrFS-TRUST Method: TTrFS Multi-normalization Multi-distance Assessment</h1>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.write("Enhanced version with soft/hard criteria handling and expert aggregation")
    
    # Initialize session state for TRUST
    if 'trust_step' not in st.session_state:
        st.session_state.trust_step = 1
    if 'trust_data' not in st.session_state:
        st.session_state.trust_data = {}
    
    # Step navigation
    steps = [
        "Problem Setup", "Criteria Setup", "Expert Weights", "Data Collection",
        "Build Decision Matrix", "Criteria Information", "Constraint Values", "Results"
    ]
    
    current_step = st.session_state.trust_step - 1
    
    # Progress bar
    progress = st.progress(current_step / (len(steps) - 1))
    st.write(f"**Current Step: {steps[current_step]}**")
    
    # Step 1: Input parameters
    if st.session_state.trust_step == 1:
        trust_step1_input()
    
    # Step 2: Criteria setup
    elif st.session_state.trust_step == 2:
        trust_step2_criteria_setup()
    
    # Step 3: Expert weights
    elif st.session_state.trust_step == 3:
        trust_step3_expert_weights()
    
    # Step 4: Data collection
    elif st.session_state.trust_step == 4:
        trust_step4_data_collection()
    
    # Step 5: Decision matrix
    elif st.session_state.trust_step == 5:
        trust_step5_decision_matrix()
    
    # Step 6: Criteria information
    elif st.session_state.trust_step == 6:
        trust_step6_criteria_info()
    
    # Step 7: Constraint values
    elif st.session_state.trust_step == 7:
        trust_step7_constraints()
    
    # Step 8: Calculations and results
    elif st.session_state.trust_step == 8:
        trust_step8_calculations()

def trust_step1_input():
    st.header("Step 1: Problem Setup")
    
    # Number of alternatives, criteria, and experts
    n_alternatives = st.number_input("Number of alternatives", min_value=2, max_value=20, value=4)
    n_criteria = st.number_input("Number of criteria", min_value=2, max_value=20, value=6)
    n_experts = st.number_input("Number of experts", min_value=1, max_value=10, value=3)
    
    # Alpha parameters for normalization
    st.subheader("Normalization Parameters (∂)")
    st.write("These parameters determine the weight of each normalization technique")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        alpha1 = st.number_input("∂₁ (Linear Ratio)", min_value=0.0, max_value=1.0, value=0.25, step=0.05)
    with col2:
        alpha2 = st.number_input("∂₂ (Linear Sum)", min_value=0.0, max_value=1.0, value=0.25, step=0.05)
    with col3:
        alpha3 = st.number_input("∂₃ (Max-Min)", min_value=0.0, max_value=1.0, value=0.25, step=0.05)
    with col4:
        alpha4 = st.number_input("∂₄ (Logarithmic)", min_value=0.0, max_value=1.0, value=0.25, step=0.05)
    
    # Check if alpha values sum to 1
    alpha_sum = alpha1 + alpha2 + alpha3 + alpha4
    if abs(alpha_sum - 1.0) > 0.001:
        st.warning(f"Alpha values sum to {alpha_sum:.2f}, but should sum to 1.0")
    
    # Beta parameter for distance aggregation
    beta = st.slider("β (Distance Aggregation Parameter)", min_value=0.0, max_value=1.0, value=0.5, step=0.1)
    
    if st.button("Next: Criteria Setup"):
        # Store parameters
        st.session_state.trust_data['n_alternatives'] = n_alternatives
        st.session_state.trust_data['n_criteria'] = n_criteria
        st.session_state.trust_data['n_experts'] = n_experts
        st.session_state.trust_data['alpha'] = [alpha1, alpha2, alpha3, alpha4]
        st.session_state.trust_data['beta'] = beta
        
        # Initialize alternatives and criteria names
        st.session_state.trust_data['alternatives'] = [f"A{i+1}" for i in range(n_alternatives)]
        st.session_state.trust_data['criteria'] = [f"C{i+1}" for i in range(n_criteria)]
        
        # Initialize criteria types and expert weights
        st.session_state.trust_data['criteria_types'] = ['Soft'] * n_criteria
        st.session_state.trust_data['expert_weights'] = [1.0/n_experts] * n_experts
        
        st.session_state.trust_step = 2
        st.rerun()

def trust_step2_criteria_setup():
    st.header("Step 2: Criteria Setup")
    
    n_criteria = st.session_state.trust_data['n_criteria']
    criteria = st.session_state.trust_data['criteria']
    
    st.write("Specify the type of each criterion (Soft or Hard) and provide descriptions")
    
    # Create dataframe for criteria setup
    data = {
        'Criterion': criteria,
        'Type': st.session_state.trust_data['criteria_types'],
        'Description': [''] * n_criteria
    }
    df = pd.DataFrame(data)
    
    # Editable table for criteria setup
    edited_df = st.data_editor(df, use_container_width=True)
    
    # Display linguistic scale information
    with st.expander("Linguistic Scale for Soft Criteria"):
        st.write("""
        **Linguistic Terms and Trapezoidal Fuzzy Numbers:**
        
        - ELI (Extremely Low Importance): (0.00, 0.10, 0.20, 0.30)
        - VLI (Very Low Importance): (0.10, 0.20, 0.30, 0.40)
        - LI (Low Importance): (0.20, 0.30, 0.40, 0.50)
        - MLI (Medium Low Importance): (0.30, 0.40, 0.50, 0.60)
        - MI (Medium Importance): (0.40, 0.50, 0.60, 0.70)
        - MHI (Medium High Importance): (0.50, 0.60, 0.70, 0.80)
        - HI (High Importance): (0.60, 0.70, 0.80, 0.90)
        - VHI (Very High Importance): (0.70, 0.80, 0.90, 1.00)
        - EHI (Extremely High Importance): (0.80, 0.90, 1.00, 1.00)
        """)
    
    if st.button("Next: Expert Weights"):
        st.session_state.trust_data['criteria_setup'] = edited_df
        st.session_state.trust_data['criteria_types'] = edited_df['Type'].values
        st.session_state.trust_step = 3
        st.rerun()

def trust_step3_expert_weights():
    st.header("Step 3: Expert Weights")
    
    n_experts = st.session_state.trust_data['n_experts']
    
    st.write("Assign weights to each expert (must sum to 1.0)")
    
    expert_weights = []
    for e in range(n_experts):
        weight = st.number_input(
            f"Weight for Expert {e+1}",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.trust_data['expert_weights'][e],
            step=0.01,
            format="%.2f"
        )
        expert_weights.append(weight)
    
    # Check if weights sum to 1
    weight_sum = sum(expert_weights)
    if abs(weight_sum - 1.0) > 0.001:
        st.warning(f"Expert weights sum to {weight_sum:.2f}, but should sum to 1.0")
        valid_weights = False
    else:
        st.success("Expert weights are valid!")
        valid_weights = True
    
    if st.button("Next: Data Collection") and valid_weights:
        st.session_state.trust_data['expert_weights'] = expert_weights
        st.session_state.trust_step = 4
        st.rerun()

def trust_step4_data_collection():
    st.header("Step 4: Data Collection")
    
    n_alternatives = st.session_state.trust_data['n_alternatives']
    n_criteria = st.session_state.trust_data['n_criteria']
    n_experts = st.session_state.trust_data['n_experts']
    alternatives = st.session_state.trust_data['alternatives']
    criteria = st.session_state.trust_data['criteria']
    criteria_types = st.session_state.trust_data['criteria_types']
    
    # Initialize expert data if not exists
    if 'expert_data' not in st.session_state.trust_data:
        st.session_state.trust_data['expert_data'] = {}
        for e in range(n_experts):
            st.session_state.trust_data['expert_data'][f'expert_{e}'] = pd.DataFrame(
                [['MI'] * n_criteria for _ in range(n_alternatives)],
                columns=criteria,
                index=alternatives
            )
    
    # Initialize hard data if not exists
    if 'hard_data' not in st.session_state.trust_data:
        st.session_state.trust_data['hard_data'] = pd.DataFrame(
            [[0.0] * n_criteria for _ in range(n_alternatives)],
            columns=criteria,
            index=alternatives
        )
    
    # Separate soft and hard criteria
    soft_criteria = [criteria[i] for i, t in enumerate(criteria_types) if t == 'Soft']
    hard_criteria = [criteria[i] for i, t in enumerate(criteria_types) if t == 'Hard']
    
    st.write("### Soft Criteria Assessment")
    st.write("Multiple experts provide linguistic assessments for soft criteria")
    
    # Expert tabs for soft criteria
    if soft_criteria:
        expert_tabs = st.tabs([f"Expert {e+1}" for e in range(n_experts)])
        
        for e, tab in enumerate(expert_tabs):
            with tab:
                st.write(f"**Expert {e+1} Assessments for Soft Criteria**")
                
                # Get current data for this expert
                current_data = st.session_state.trust_data['expert_data'][f'expert_{e}']
                
                # Filter only soft criteria
                soft_data = current_data[soft_criteria].copy()
                
                # Data editor for soft criteria
                edited_soft_data = st.data_editor(
                    soft_data,
                    column_config={
                        col: st.column_config.SelectboxColumn(
                            col,
                            options=linguistic_options_trust,
                            required=True
                        ) for col in soft_criteria
                    },
                    use_container_width=True,
                    key=f"trust_expert_{e}_soft"
                )
                
                # Update session state
                st.session_state.trust_data['expert_data'][f'expert_{e}'].update(edited_soft_data)
    
    st.write("### Hard Criteria Assessment")
    st.write("Enter crisp values for hard criteria (single value per alternative)")
    
    if hard_criteria:
        # Data editor for hard criteria
        current_hard_data = st.session_state.trust_data['hard_data'][hard_criteria].copy()
        edited_hard_data = st.data_editor(
            current_hard_data,
            use_container_width=True,
            key="trust_hard_data"
        )
        
        # Update session state
        st.session_state.trust_data['hard_data'].update(edited_hard_data)
    
    if st.button("Next: Build Decision Matrix"):
        st.session_state.trust_step = 5
        st.rerun()

def trust_step5_decision_matrix():
    st.header("Step 5: Build Decision Matrix")
    
    n_alternatives = st.session_state.trust_data['n_alternatives']
    n_criteria = st.session_state.trust_data['n_criteria']
    n_experts = st.session_state.trust_data['n_experts']
    alternatives = st.session_state.trust_data['alternatives']
    criteria = st.session_state.trust_data['criteria']
    criteria_types = st.session_state.trust_data['criteria_types']
    expert_weights = st.session_state.trust_data['expert_weights']
    
    st.write("Aggregating expert assessments and building final decision matrix...")
    
    # Initialize decision matrix
    decision_matrix = np.zeros((n_alternatives, n_criteria))
    
    # Initialize matrices to store intermediate results for display
    expert_assessments_matrix = [[[] for _ in range(n_criteria)] for _ in range(n_alternatives)]
    aggregated_trfn_matrix = [[() for _ in range(n_criteria)] for _ in range(n_alternatives)]
    
    # Process each criterion
    for j, criterion in enumerate(criteria):
        if criteria_types[j] == 'Soft':
            # Collect all expert assessments for this criterion
            all_expert_assessments = []
            for i in range(n_alternatives):
                alternative_assessments = []
                for e in range(n_experts):
                    expert_df = st.session_state.trust_data['expert_data'][f'expert_{e}']
                    linguistic_value = expert_df.loc[alternatives[i], criterion]
                    trfn_value = ling_to_trfn_trust[linguistic_value]
                    alternative_assessments.append(trfn_value)
                all_expert_assessments.append(alternative_assessments)
            
            # Aggregate using trigonometric trapezoidal fuzzy weighted geometric (TTrFWG) for each alternative
            for i in range(n_alternatives):
                aggregated_trfn = aggregate_ftwg(all_expert_assessments[i], expert_weights)
                defuzzified_value = defuzzify_trfn(aggregated_trfn)
                decision_matrix[i, j] = defuzzified_value
                
                # Store intermediate results for display
                expert_assessments_matrix[i][j] = [st.session_state.trust_data['expert_data'][f'expert_{e}'].loc[alternatives[i], criterion] for e in range(n_experts)]
                aggregated_trfn_matrix[i][j] = aggregated_trfn
        
        else:  # Hard criterion
            for i in range(n_alternatives):
                crisp_value = st.session_state.trust_data['hard_data'].loc[alternatives[i], criterion]
                trfn_value = crisp_to_trfs(crisp_value)
                defuzzified_value = defuzzify_trfn(trfn_value)
                decision_matrix[i, j] = defuzzified_value
                
                # Store intermediate results for display
                expert_assessments_matrix[i][j] = [crisp_value]  # Store crisp value for hard criteria
                aggregated_trfn_matrix[i][j] = trfn_value
    
    # Display Aggregated Matrix in a clean tabular format
    st.subheader("Aggregated Decision Matrix Details")
    
    # Create tabs for each alternative
    alt_tabs = st.tabs(alternatives)
    
    for tab_idx, (tab, alternative) in enumerate(zip(alt_tabs, alternatives)):
        with tab:
            st.write(f"**{alternative} - Aggregated Assessments**")
            
            # Create data for the table
            table_data = []
            for j, criterion in enumerate(criteria):
                expert_assessments = expert_assessments_matrix[tab_idx][j]
                aggregated_trfn = aggregated_trfn_matrix[tab_idx][j]
                defuzzified_value = decision_matrix[tab_idx, j]
                
                # Format expert assessments
                if criteria_types[j] == 'Soft':
                    assessments_str = ", ".join(expert_assessments)
                else:
                    assessments_str = f"Crisp: {expert_assessments[0]}"
                
                # Format aggregated TrFN
                trfn_str = f"({aggregated_trfn[0]:.4f}, {aggregated_trfn[1]:.4f}, {aggregated_trfn[2]:.4f}, {aggregated_trfn[3]:.4f})"
                
                table_data.append({
                    'Criterion': criterion,
                    'Type': criteria_types[j],
                    'Expert Assessments': assessments_str,
                    'Aggregated TrFN': trfn_str,
                    'Defuzzified Value': f"{defuzzified_value:.4f}"
                })
            
            # Display as dataframe
            df_detailed = pd.DataFrame(table_data)
            st.dataframe(df_detailed, use_container_width=True, hide_index=True)
    
    # Display final decision matrix
    st.subheader("Final Decision Matrix (Defuzzified Values)")
    df_decision = pd.DataFrame(decision_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_decision, use_container_width=True)
    
    # Display summary statistics
    st.subheader("Summary Statistics")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Total Alternatives", n_alternatives)
    with col2:
        st.metric("Total Criteria", n_criteria)
    with col3:
        soft_count = sum(1 for ct in criteria_types if ct == 'Soft')
        hard_count = n_criteria - soft_count
        st.metric("Soft/Hard Criteria", f"{soft_count}/{hard_count}")
    
    if st.button("Next: Criteria Information"):
        st.session_state.trust_data['decision_matrix'] = decision_matrix
        st.session_state.trust_step = 6
        st.rerun()

def trust_step6_criteria_info():
    st.header("Step 6: Criteria Information")
    
    n_criteria = st.session_state.trust_data['n_criteria']
    criteria = st.session_state.trust_data['criteria']
    
    st.write("Specify the type of each criterion (Benefit or Cost) and their weights")
    
    # Create dataframe for criteria information
    data = {
        'Criterion': criteria,
        'Type': ['Benefit'] * n_criteria,
        'Weight': [1.0/n_criteria] * n_criteria
    }
    df = pd.DataFrame(data)
    
    # Editable table for criteria information
    edited_df = st.data_editor(df, use_container_width=True)
    
    # Check if weights sum to 1
    weight_sum = edited_df['Weight'].sum()
    if abs(weight_sum - 1.0) > 0.001:
        st.warning(f"Weights sum to {weight_sum:.2f}, but should sum to 1.0")
    
    if st.button("Next: Constraint Values"):
        st.session_state.trust_data['criteria_info'] = edited_df
        st.session_state.trust_step = 7
        st.rerun()

def trust_step7_constraints():
    st.header("Step 7: Constraint Values")
    
    n_criteria = st.session_state.trust_data['n_criteria']
    criteria = st.session_state.trust_data['criteria']
    decision_matrix = st.session_state.trust_data['decision_matrix']
    
    st.write("Specify the constraint intervals for each criterion")
    st.write("ϱⱼᴸ: Lower Bound, ϱⱼᵁ: Upper Bound")
    
    # Calculate min and max values for each criterion
    min_vals = np.min(decision_matrix, axis=0)
    max_vals = np.max(decision_matrix, axis=0)
    
    # Create dataframe for constraint values
    data = {
        'Criterion': criteria,
        'Min Value': min_vals,
        'Max Value': max_vals,
        'ϱⱼᴸ': min_vals,
        'ϱⱼᵁ': max_vals
    }
    df = pd.DataFrame(data)
    
    # Editable table for constraint values
    edited_df = st.data_editor(df, use_container_width=True)
    
    if st.button("Calculate Results"):
        st.session_state.trust_data['constraints'] = edited_df
        st.session_state.trust_step = 8
        st.rerun()

def trust_step8_calculations():
    st.header("Step 8: TRUST Method Results")
    
    # Retrieve data from session state
    alternatives = st.session_state.trust_data['alternatives']
    criteria = st.session_state.trust_data['criteria']
    decision_matrix = st.session_state.trust_data['decision_matrix']
    criteria_info = st.session_state.trust_data['criteria_info']
    constraints = st.session_state.trust_data['constraints']
    alpha = st.session_state.trust_data['alpha']
    beta = st.session_state.trust_data['beta']
    
    n_alternatives = len(alternatives)
    n_criteria = len(criteria)
    
    # Extract criteria types and weights
    criteria_types = criteria_info['Type'].values
    weights = criteria_info['Weight'].values
    
    # Extract constraint values
    LB = constraints['ϱⱼᴸ'].values
    UB = constraints['ϱⱼᵁ'].values
    
    # Store all data for Word export
    all_data = {
        'n_alternatives': n_alternatives,
        'n_criteria': n_criteria,
        'n_experts': st.session_state.trust_data['n_experts'],
        'alpha': alpha,
        'beta': beta,
        'alternatives': alternatives,
        'criteria': criteria,
        'decision_matrix': decision_matrix
    }
    
    # Step 2.3: Normalization
    st.subheader("Step 2.3: Normalization")
    
    # Calculate min and max for each criterion
    min_vals = np.min(decision_matrix, axis=0)
    max_vals = np.max(decision_matrix, axis=0)
    
    # Normalization type 1: Linear ratio-based (r_ij)
    st.write("Normalization Type 1: Linear Ratio-based (r_ij)")
    r_matrix = np.zeros((n_alternatives, n_criteria))
    for j in range(n_criteria):
        if criteria_types[j] == 'Benefit':
            max_val = np.max(decision_matrix[:, j])
            r_matrix[:, j] = decision_matrix[:, j] / max_val
        else:  # Cost criterion
            min_val = np.min(decision_matrix[:, j])
            r_matrix[:, j] = min_val / decision_matrix[:, j]
    
    df_r = pd.DataFrame(r_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_r, use_container_width=True)
    all_data['r_matrix'] = r_matrix
    
    # Normalization type 2: Linear sum-based (s_ij)
    st.write("Normalization Type 2: Linear Sum-based (s_ij)")
    s_matrix = np.zeros((n_alternatives, n_criteria))
    for j in range(n_criteria):
        if criteria_types[j] == 'Benefit':
            sum_val = np.sum(decision_matrix[:, j])
            s_matrix[:, j] = decision_matrix[:, j] / sum_val
        else:  # Cost criterion
            sum_reciprocal = np.sum(1 / decision_matrix[:, j])
            s_matrix[:, j] = (1 / decision_matrix[:, j]) / sum_reciprocal
    
    df_s = pd.DataFrame(s_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_s, use_container_width=True)
    all_data['s_matrix'] = s_matrix
    
    # Normalization type 3: Linear max-min (m_ij)
    st.write("Normalization Type 3: Linear Max-Min (m_ij)")
    m_matrix = np.zeros((n_alternatives, n_criteria))
    for j in range(n_criteria):
        if criteria_types[j] == 'Benefit':
            m_matrix[:, j] = (decision_matrix[:, j] - min_vals[j]) / (max_vals[j] - min_vals[j])
        else:  # Cost criterion
            m_matrix[:, j] = (max_vals[j] - decision_matrix[:, j]) / (max_vals[j] - min_vals[j])
    
    df_m = pd.DataFrame(m_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_m, use_container_width=True)
    all_data['m_matrix'] = m_matrix
    
    # Normalization type 4: Logarithmic (l_ij)
    st.write("Normalization Type 4: Logarithmic (l_ij)")
    l_matrix = np.zeros((n_alternatives, n_criteria))
    for j in range(n_criteria):
        product = np.prod(decision_matrix[:, j])
        for i in range(n_alternatives):
            if decision_matrix[i, j] > 0:
                l_matrix[i, j] = np.log(decision_matrix[i, j]) / np.log(product)
    
    df_l = pd.DataFrame(l_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_l, use_container_width=True)
    all_data['l_matrix'] = l_matrix
    
    # Step 2.3.5: Aggregate normalized matrices (h_ij)
    st.write("Aggregated Normalized Matrix (h_ij)")
    h_matrix = alpha[0] * r_matrix + alpha[1] * s_matrix + alpha[2] * m_matrix + alpha[3] * l_matrix
    df_h = pd.DataFrame(h_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_h, use_container_width=True)
    all_data['h_matrix'] = h_matrix
    
    # Step 2.4: Constraint-based normalization
    st.subheader("Step 2.4: Constraint-based Normalization")
    
    # Calculate satisfaction degree matrix (f_ij)
    f_matrix = np.zeros((n_alternatives, n_criteria))
    for j in range(n_criteria):
        for i in range(n_alternatives):
            d_ij = decision_matrix[i, j]
            lb_j = LB[j]
            ub_j = UB[j]
            min_j = min_vals[j]
            max_j = max_vals[j]
            max_denom = max(lb_j - min_j, max_j - ub_j)
            
            if criteria_types[j] == 'Benefit':
                if lb_j <= d_ij <= ub_j:
                    f_matrix[i, j] = 1.0
                elif d_ij < lb_j:
                    f_matrix[i, j] = 1 - (lb_j - d_ij) / (max_denom + 1)
                else:  # d_ij > ub_j
                    f_matrix[i, j] = 1 - (1 - ub_j + d_ij) / (max_denom + 1)
            else:  # Cost criterion
                if lb_j <= d_ij <= ub_j:
                    f_matrix[i, j] = 1 / (max_denom + 1)
                elif d_ij < lb_j:
                    f_matrix[i, j] = (lb_j - d_ij) / max_denom
                else:  # d_ij > ub_j
                    f_matrix[i, j] = (d_ij - ub_j) / max_denom
    
    df_f = pd.DataFrame(f_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_f, use_container_width=True)
    all_data['f_matrix'] = f_matrix
    
    # Step 2.5: Constrained aggregated score matrix (η_ij)
    st.subheader("Step 2.5: Constrained Aggregated Score Matrix (η_ij)")
    eta_matrix = h_matrix * f_matrix
    df_eta = pd.DataFrame(eta_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_eta, use_container_width=True)
    all_data['eta_matrix'] = eta_matrix
    
    # Step 2.6: Weighted decision matrix (v_ij)
    st.subheader("Step 2.6: Weighted Decision Matrix (v_ij)")
    v_matrix = eta_matrix * weights
    df_v = pd.DataFrame(v_matrix, columns=criteria, index=alternatives)
    st.dataframe(df_v, use_container_width=True)
    all_data['v_matrix'] = v_matrix
    
    # Step 2.7: Negative-ideal solution (τ_j)
    st.subheader("Step 2.7: Negative-Ideal Solution (τ_j)")
    tau = np.min(v_matrix, axis=0)
    df_tau = pd.DataFrame([tau], columns=criteria, index=['τ'])
    st.dataframe(df_tau, use_container_width=True)
    all_data['tau'] = tau
    
    # Step 2.8: Distance measures
    st.subheader("Step 2.8: Distance Measures")
    
    # Euclidean distance (ε_i)
    epsilon = np.sqrt(np.sum((v_matrix - tau) ** 2, axis=1))
    
    # Manhattan distance (π_i)
    pi = np.sum(np.abs(v_matrix - tau), axis=1)
    
    # Lorentzian distance (l_i)
    st.write("Lorentzian Distance Calculation")
    l_distance = np.zeros(n_alternatives)
    for i in range(n_alternatives):
        lorentz_sum = 0
        for j in range(n_criteria):
            diff = abs(v_matrix[i, j] - tau[j])
            lorentz_sum += np.log10(1 + diff)
        l_distance[i] = lorentz_sum
    
    # Display detailed Lorentzian calculation
    st.write("Detailed Lorentzian Distance Calculation:")
    lorentz_details = []
    for i in range(n_alternatives):
        row_details = []
        for j in range(n_criteria):
            diff = abs(v_matrix[i, j] - tau[j])
            log_val = np.log10(1 + diff)
            row_details.append(f"{log_val:.4f}")
        lorentz_details.append(f"Sum({', '.join(row_details)})")
    
    # Create a DataFrame for better visualization
    lorentz_df = pd.DataFrame({
        'Alternative': alternatives,
        'Lorentzian Distance': l_distance,
        'Calculation Details': lorentz_details
    })
    st.dataframe(lorentz_df, use_container_width=True)
    all_data['l_distance'] = l_distance
    
    # Pearson distance (ρ_i)
    tau_safe = np.where(tau == 0, 1e-9, tau)
    rho = np.sum(((v_matrix - tau) ** 2) / tau_safe, axis=1)
    all_data['rho'] = rho
    
    # Create dataframe for distance measures
    distances = pd.DataFrame({
        'Euclidean (ε)': epsilon,
        'Manhattan (π)': pi,
        'Lorentzian (l)': l_distance,
        'Pearson (ρ)': rho
    }, index=alternatives)
    st.dataframe(distances, use_container_width=True)
    all_data['distances'] = distances
    
    # Step 2.9: Relative assessment matrices
    st.subheader("Step 2.9: Relative Assessment Matrices")
    
    # Calculate relative assessment matrices
    wp_sum = np.zeros(n_alternatives)  # ∑℘_ik
    H_sum = np.zeros(n_alternatives)  # ∑ℌ_ik
    
    for i in range(n_alternatives):
        for k in range(n_alternatives):
            # ℘_ik = (ε_i - ε_k) + ((ε_i - ε_k) × (π_i - π_k))
            wp_ik = (epsilon[i] - epsilon[k]) + ((epsilon[i] - epsilon[k]) * (pi[i] - pi[k]))
            wp_sum[i] += wp_ik
            
            # ℌ_ik = (l_i - l_k) + ((l_i - l_k) × (ρ_i - ρ_k))
            H_ik = (l_distance[i] - l_distance[k]) + ((l_distance[i] - l_distance[k]) * (rho[i] - rho[k]))
            H_sum[i] += H_ik
    
    # Display relative assessment matrices
    rel_assessment = pd.DataFrame({
        'Alternative': alternatives,
        '∑℘_ik': wp_sum,
        '∑ℌ_ik': H_sum
    })
    st.dataframe(rel_assessment, use_container_width=True)
    all_data['rel_assessment'] = rel_assessment
    
    # Step 2.10: Combined score and ranking
    st.subheader("Step 2.10: Final Scores and Ranking")
    
    # Calculate combined score (ℒ_i)
    L_score = beta * wp_sum + (1 - beta) * H_sum
    
    # Create results dataframe
    results = pd.DataFrame({
        'Alternative': alternatives,
        '∑℘_ik': wp_sum,
        '∑ℌ_ik': H_sum,
        'ℒ Score': L_score
    })
    
    # Sort by score in descending order
    results = results.sort_values('ℒ Score', ascending=False)
    results['Rank'] = range(1, len(results) + 1)
    
    st.dataframe(results, use_container_width=True)
    
    # Display best alternative
    best_alt = results.iloc[0]['Alternative']
    best_score = results.iloc[0]['ℒ Score']
    st.success(f"**Best Alternative: {best_alt}** with score: {best_score:.4f}")
    
    all_data['final_results'] = results
    all_data['best_alternative'] = best_alt
    all_data['best_score'] = best_score
    
    # Export button
    st.markdown('<div style="text-align: center; margin: 2.5rem 0;">', unsafe_allow_html=True)
    doc_bytes = create_trust_word_document(all_data)
    
    st.download_button(
        label="Export TRUST Results to Word",
        data=doc_bytes,
        file_name="TRUST_Method_Results.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Option to restart
    if st.button("Start Over"):
        st.session_state.trust_step = 1
        st.session_state.trust_data = {}
        st.rerun()

def trust_model():
    with st.expander("Learn more about the Trigonometric Trapezoidal Fuzzy TRUST (TTrFS-TRUST) Method"):
        st.markdown("""
        **Overview of Trigonometric Trapezoidal Fuzzy TRUST (TTrFS-TRUST) Method:**
        
        The TTrFS-TRUST (Trigonometric Trapezoidal Fuzzy Soft Multi-normalization Multi-distance Assessment) method 
        is an advanced MCDM approach that combines trigonometric trapezoidal fuzzy sets with multiple normalization 
        techniques and distance measures for comprehensive alternative evaluation.
        
        **Key Steps of TTrFS-TRUST Method:**
        
        1. **Problem Setup**: Define alternatives, criteria, experts, and parameters (∂ and β)
        2. **Criteria Classification**: Separate criteria into Soft (linguistic assessments) and Hard (crisp values)
        3. **Expert Aggregation**: Use Trigonometric Trapezoidal Fuzzy Weighted Geometric (TTrFWG) operator to aggregate multiple expert opinions
        4. **Multi-Normalization**: Apply four different normalization techniques:
           - Linear Ratio-based (r_ij)
           - Linear Sum-based (s_ij) 
           - Max-Min (m_ij)
           - Logarithmic (l_ij)
        5. **Constraint Integration**: Incorporate constraint intervals (ϱⱼᴸ, ϱⱼᵁ) for realistic boundary conditions
        6. **Multi-Distance Assessment**: Calculate four distance measures from negative-ideal solution:
           - Euclidean Distance (ε_i)
           - Manhattan Distance (π_i)
           - Lorentzian Distance (l_i)
           - Pearson Distance (ρ_i)
        7. **Relative Assessment**: Compute relative assessment matrices (℘_ik and ℌ_ik)
        8. **Final Ranking**: Combine distances using β parameter to obtain final scores (ℒ_i)
        
        **Mathematical Formulation:**
        
        **Aggregated Normalized Matrix:**
        h_ij = ∂₁·r_ij + ∂₂·s_ij + ∂₃·m_ij + ∂₄·l_ij
        
        **Constrained Score Matrix:**
        η_ij = h_ij × f_ij  (where f_ij is constraint satisfaction degree)
        
        **Weighted Decision Matrix:**
        v_ij = η_ij × w_j
        
        **Distance Measures:**
        - Euclidean: ε_i = √[∑(v_ij - τ_j)²]
        - Manhattan: π_i = ∑|v_ij - τ_j|
        - Lorentzian: l_i = ∑log₁₀(1 + |v_ij - τ_j|)
        - Pearson: ρ_i = ∑[(v_ij - τ_j)² / τ_j]
        
        **Relative Assessment Matrices:**
        ℘_ik = (ε_i - ε_k) + (ε_i - ε_k) × (π_i - π_k)
        ℌ_ik = (l_i - l_k) + (l_i - l_k) × (ρ_i - ρ_k)
        
        **Final Score:**
        ℒ_i = β·∑℘_ik + (1-β)·∑ℌ_ik
        
        **Linguistic Scale for Soft Criteria:**
        - ELI (Extremely Low): (0.00, 0.10, 0.20, 0.30)
        - VLI (Very Low): (0.10, 0.20, 0.30, 0.40)
        - LI (Low): (0.20, 0.30, 0.40, 0.50)
        - MLI (Medium Low): (0.30, 0.40, 0.50, 0.60)
        - MI (Medium): (0.40, 0.50, 0.60, 0.70)
        - MHI (Medium High): (0.50, 0.60, 0.70, 0.80)
        - HI (High): (0.60, 0.70, 0.80, 0.90)
        - VHI (Very High): (0.70, 0.80, 0.90, 1.00)
        - EHI (Extremely High): (0.80, 0.90, 1.00, 1.00)
        
        **Parameters:**
        - **∂ = (∂₁, ∂₂, ∂₃, ∂₄)**: Normalization weights (must sum to 1)
        - **β**: Distance aggregation parameter (0 ≤ β ≤ 1)
        - **ϱⱼᴸ, ϱⱼᵁ**: Lower and upper constraint bounds for criterion j
        
        **Advantages of TTrFS-TRUST:**
        - Handles both linguistic and crisp data simultaneously
        - Incorporates multiple expert opinions with weighted aggregation
        - Uses multiple normalization techniques for robust results
        - Considers constraint-based realistic boundaries
        - Combines multiple distance measures for comprehensive assessment
        - Provides transparent step-by-step calculation process
        
        **Applications:**
        - Complex decision-making problems with mixed data types
        - Situations requiring multiple expert input
        - Problems with constraint boundaries and realistic limits
        - Cases where robust, multi-perspective evaluation is needed
        """)

# ==================== MAIN APP ====================

def main():
    st.sidebar.title("TTrFS MCDM Model Selection")
    st.sidebar.markdown("Select the model you want to use:")
    
    model_choice = st.sidebar.radio(
        "Choose Model:",
        ["Trigonometric Trapezoidal Fuzzy OPA", "TTrFS-TRUST Method"],
        index=0
    )
    
    if model_choice == "Trigonometric Trapezoidal Fuzzy OPA":
        opa_model()
    elif model_choice == "TTrFS-TRUST Method":
        trust_model()
    
    st.markdown("""
    <div class="footer">
    <p>Integrated MCDM Models Implementation | Developed by AAA</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
